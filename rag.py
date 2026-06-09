"""
Local RAG over your own error notebooks & reference materials.

Drop files into ./knowledge_base/ (txt, md, csv, pdf, docx, xlsx). On first run we
chunk + embed them with a small local model (sentence-transformers, free, no API
cost) and cache the vectors to ./knowledge_base/.kb_cache.npz. Retrieval is plain
cosine similarity in numpy — fast and dependency-light for a personal KB.

Nothing here calls a paid API. The embeddings run on your machine / the Streamlit
host. Only the final, retrieved snippets get sent to Groq as context.
"""

import os
import glob
import hashlib
import numpy as np

KB_DIR = os.path.join(os.path.dirname(__file__), "knowledge_base")
CACHE = os.path.join(KB_DIR, ".kb_cache.npz")
EMBED_MODEL = "all-MiniLM-L6-v2"   # ~80MB, downloaded once, runs locally
CHUNK_CHARS = 900
CHUNK_OVERLAP = 150

_model = None


def _get_model():
    global _model
    if _model is None:
        from sentence_transformers import SentenceTransformer
        _model = SentenceTransformer(EMBED_MODEL)
    return _model


# ---------- file loaders (graceful: skip a format if its lib is missing) ----------

def _read_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _read_pdf(path):
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        reader = PdfReader(path)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        return ""


def _read_docx(path):
    try:
        import docx
    except ImportError:
        return ""
    try:
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception:
        return ""


def _read_tabular(path):
    try:
        import pandas as pd
    except ImportError:
        return ""
    try:
        if path.lower().endswith((".xlsx", ".xls")):
            dfs = pd.read_excel(path, sheet_name=None)
            return "\n\n".join(f"# sheet: {n}\n{df.to_csv(index=False)}" for n, df in dfs.items())
        return pd.read_csv(path).to_csv(index=False)
    except Exception:
        return ""


LOADERS = {
    ".txt": _read_txt, ".md": _read_txt, ".log": _read_txt,
    ".pdf": _read_pdf, ".docx": _read_docx,
    ".csv": _read_tabular, ".xlsx": _read_tabular, ".xls": _read_tabular,
}


def _chunk(text, source):
    text = text.strip()
    chunks = []
    i = 0
    while i < len(text):
        piece = text[i:i + CHUNK_CHARS].strip()
        if piece:
            chunks.append({"text": piece, "source": source})
        i += CHUNK_CHARS - CHUNK_OVERLAP
    return chunks


def _collect_files():
    files = []
    for ext in LOADERS:
        files.extend(glob.glob(os.path.join(KB_DIR, "**", f"*{ext}"), recursive=True))
    return sorted(f for f in files if not os.path.basename(f).startswith("."))


def _fingerprint(files):
    """Hash of (path, mtime, size) so we re-embed only when files change."""
    h = hashlib.sha256()
    for f in files:
        st = os.stat(f)
        h.update(f.encode()); h.update(str(st.st_mtime_ns).encode()); h.update(str(st.st_size).encode())
    return h.hexdigest()


# ---------------------------- public API ----------------------------

def build_index(force=False):
    """Build (or load cached) the vector index. Returns (n_chunks, n_files)."""
    files = _collect_files()
    if not files:
        return 0, 0

    fp = _fingerprint(files)
    if not force and os.path.exists(CACHE):
        cached = np.load(CACHE, allow_pickle=True)
        if str(cached.get("fingerprint")) == fp:
            return int(cached["vectors"].shape[0]), len(files)

    all_chunks = []
    for f in files:
        ext = os.path.splitext(f)[1].lower()
        text = LOADERS[ext](f)
        if text and text.strip():
            all_chunks.extend(_chunk(text, os.path.basename(f)))

    if not all_chunks:
        return 0, len(files)

    model = _get_model()
    texts = [c["text"] for c in all_chunks]
    vectors = model.encode(texts, normalize_embeddings=True, show_progress_bar=False)
    np.savez(
        CACHE,
        vectors=np.asarray(vectors, dtype=np.float32),
        texts=np.array(texts, dtype=object),
        sources=np.array([c["source"] for c in all_chunks], dtype=object),
        fingerprint=fp,
    )
    return len(all_chunks), len(files)


def retrieve(query, k=4, min_score=0.25):
    """Return up to k relevant chunks as list of {text, source, score}."""
    if not os.path.exists(CACHE):
        return []
    cached = np.load(CACHE, allow_pickle=True)
    vectors = cached["vectors"]
    if vectors.shape[0] == 0:
        return []
    qv = _get_model().encode([query], normalize_embeddings=True)[0]
    scores = vectors @ qv                       # cosine sim (vectors are normalized)
    top = np.argsort(scores)[::-1][:k]
    out = []
    for i in top:
        if float(scores[i]) >= min_score:
            out.append({"text": str(cached["texts"][i]),
                        "source": str(cached["sources"][i]),
                        "score": round(float(scores[i]), 3)})
    return out


def context_block(query, k=4):
    """Format retrieved chunks as a context string for the prompt (or '')."""
    hits = retrieve(query, k=k)
    if not hits:
        return "", []
    block = "REFERENCE MATERIAL (from your knowledge base):\n\n"
    for h in hits:
        block += f"[source: {h['source']} | relevance {h['score']}]\n{h['text']}\n\n"
    return block.strip(), hits
